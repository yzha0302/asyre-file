#!/usr/bin/env python3
"""Markdown export module — PDF (Playwright) and Word (python-docx)."""

import json, os, re, uuid, time, base64, subprocess, tempfile, threading, glob
from datetime import datetime
from docx.shared import Pt, Cm

WORKSPACE = os.path.expanduser('~/xiu-he')
_EDITOR_DIR = os.path.dirname(os.path.abspath(__file__))

# ==================== CONFIG LOADER ====================
_CONFIG_PATH = os.path.join(_EDITOR_DIR, 'config.json')
_config_cache = None
_config_mtime = 0

def get_config():
    """Load config.json with mtime-based cache — hot-reloadable without restart."""
    global _config_cache, _config_mtime
    try:
        mt = os.path.getmtime(_CONFIG_PATH)
        if _config_cache is None or mt > _config_mtime:
            with open(_CONFIG_PATH, 'r') as f:
                _config_cache = json.load(f)
            _config_mtime = mt
    except:
        if _config_cache is None:
            _config_cache = {}  # fallback empty
    return _config_cache

def _resolve_avatar(avatar_rel):
    """Resolve avatar path: relative to md-editor dir first, then workspace."""
    if not avatar_rel:
        return None
    # Try relative to md-editor directory
    p = os.path.join(_EDITOR_DIR, avatar_rel)
    if os.path.isfile(p):
        return p
    # Fallback: relative to workspace
    p = os.path.join(WORKSPACE, avatar_rel)
    if os.path.isfile(p):
        return p
    return None

def _get_signatures():
    """Get signatures from config, with hardcoded defaults as fallback."""
    config = get_config()
    cfg_sigs = config.get('signatures', None)
    if cfg_sigs is None:
        # Hardcoded defaults
        return {
            'asher': {'name': 'Asher 修荷', 'avatar': os.path.join(WORKSPACE, 'assets/avatars/asher-avatar.jpg')},
            'tuisheng': {'name': '蜕升学院', 'avatar': None},
            'none': None,
        }
    result = {}
    for key, val in cfg_sigs.items():
        if val is None:
            result[key] = None
        else:
            result[key] = {
                'name': val.get('name', key),
                'avatar': _resolve_avatar(val.get('avatar')),
            }
    return result

# Backward-compatible module-level reference (reads from config each time)
class _SignaturesProxy(dict):
    """Proxy that reads from config on each access for backward compat."""
    def __getitem__(self, key):
        return _get_signatures()[key]
    def get(self, key, default=None):
        return _get_signatures().get(key, default)
    def __contains__(self, key):
        return key in _get_signatures()
    def items(self):
        return _get_signatures().items()
    def keys(self):
        return _get_signatures().keys()

SIGNATURES = _SignaturesProxy()

# ==================== CLEANUP ====================
_cleanup_started = False

def start_cleanup_thread():
    """Start background thread to clean old export files."""
    global _cleanup_started
    if _cleanup_started:
        return
    _cleanup_started = True
    def _cleanup_loop():
        while True:
            time.sleep(60)
            try:
                now = time.time()
                for f in glob.glob('/tmp/md-export-*'):
                    if now - os.path.getmtime(f) > 300:  # 5 minutes
                        os.unlink(f)
            except:
                pass
    t = threading.Thread(target=_cleanup_loop, daemon=True)
    t.start()

# ==================== MERMAID ====================
def _render_mermaid_svg(mermaid_code):
    """Render mermaid code to SVG string using mmdc."""
    fin_path = None
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as fin:
            fin.write(mermaid_code)
            fin_path = fin.name
        out_path = fin_path + '.svg'
        subprocess.run(
            ['mmdc', '-i', fin_path, '-o', out_path, '-b', 'transparent', '-q'],
            capture_output=True, timeout=30
        )
        if os.path.isfile(out_path):
            with open(out_path, 'r') as f:
                svg = f.read()
            os.unlink(out_path)
            os.unlink(fin_path)
            return svg
    except:
        pass
    finally:
        if fin_path:
            for p in [fin_path, fin_path + '.svg']:
                try: os.unlink(p)
                except: pass
    return None

def _render_mermaid_png(mermaid_code):
    """Render mermaid code to PNG bytes using mmdc."""
    fin_path = None
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as fin:
            fin.write(mermaid_code)
            fin_path = fin.name
        out_path = fin_path + '.png'
        subprocess.run(
            ['mmdc', '-i', fin_path, '-o', out_path, '-b', 'white', '-q', '-s', '2'],
            capture_output=True, timeout=30
        )
        if os.path.isfile(out_path):
            with open(out_path, 'rb') as f:
                data = f.read()
            os.unlink(out_path)
            os.unlink(fin_path)
            return data
    except:
        pass
    finally:
        if fin_path:
            for p in [fin_path, fin_path + '.png']:
                try: os.unlink(p)
                except: pass
    return None

def _render_mermaid_svgs_batch(codes):
    """Render multiple mermaid codes to SVG. Uses single mmdc call with .md input for speed."""
    if not codes:
        return []
    # Build a single markdown file with all mermaid blocks
    md_parts = []
    for i, code in enumerate(codes):
        md_parts.append(f'```mermaid\n{code}\n```\n\n---SPLIT{i}---\n')
    md_content = '\n'.join(md_parts)

    results = [None] * len(codes)
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as fin:
            fin.write(md_content)
            fin_path = fin.name
        out_path = fin_path + '.svg'  # mmdc will create -0.svg, -1.svg, etc.
        subprocess.run(
            ['mmdc', '-i', fin_path, '-o', out_path, '-b', 'transparent', '-q'],
            capture_output=True, timeout=120
        )
        # mmdc outputs: {base}-1.svg, {base}-2.svg, ... (1-based for .md input)
        base = out_path.rsplit('.svg', 1)[0]
        for i in range(len(codes)):
            svg_file = f'{base}-{i+1}.svg'
            if os.path.isfile(svg_file):
                with open(svg_file, 'r') as f:
                    results[i] = f.read()
                os.unlink(svg_file)
    except Exception:
        pass
    finally:
        try: os.unlink(fin_path)
        except: pass
    return results

def _render_mermaid_pngs_batch(codes):
    """Render multiple mermaid codes to PNG. Uses single mmdc call with .md input for speed."""
    if not codes:
        return []
    md_parts = []
    for i, code in enumerate(codes):
        md_parts.append(f'```mermaid\n{code}\n```\n')
    md_content = '\n'.join(md_parts)

    results = [None] * len(codes)
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as fin:
            fin.write(md_content)
            fin_path = fin.name
        out_path = fin_path + '.png'
        subprocess.run(
            ['mmdc', '-i', fin_path, '-o', out_path, '-b', 'white', '-q', '-s', '2'],
            capture_output=True, timeout=120
        )
        base = out_path.rsplit('.png', 1)[0]
        for i in range(len(codes)):
            png_file = f'{base}-{i+1}.png'
            if os.path.isfile(png_file):
                with open(png_file, 'rb') as f:
                    results[i] = f.read()
                os.unlink(png_file)
    except Exception:
        pass
    finally:
        try: os.unlink(fin_path)
        except: pass
    return results

# ==================== MARKDOWN → HTML ====================
def _md_to_html(md_text, theme='dark'):
    """Convert markdown to styled HTML with Pygments syntax highlighting and mermaid support."""
    import markdown
    from pygments import highlight
    from pygments.lexers import get_lexer_by_name, guess_lexer, TextLexer
    from pygments.formatters import HtmlFormatter

    # Pre-process: extract mermaid blocks, render in parallel
    mermaid_codes = []
    mermaid_keys = []
    def extract_mermaid(m):
        code = m.group(1).strip()
        key = f'XMERMAIDX{len(mermaid_codes)}X'
        mermaid_codes.append(code)
        mermaid_keys.append(key)
        return key

    md_text = re.sub(r'```mermaid\s*\n(.*?)```', extract_mermaid, md_text, flags=re.DOTALL)

    # Render all mermaid SVGs in one batch (single mmdc process)
    mermaid_blocks = {}
    if mermaid_codes:
        svgs = _render_mermaid_svgs_batch(mermaid_codes)
        for key, code, svg in zip(mermaid_keys, mermaid_codes, svgs):
            if svg:
                mermaid_blocks[key] = f'<div class="mermaid-container">{svg}</div>'
            else:
                mermaid_blocks[key] = f'<pre class="mermaid-fallback"><code>{code}</code></pre>'

    # Pre-process: syntax highlight code blocks with Pygments
    code_blocks = {}
    def replace_code(m):
        lang = m.group(1) or ''
        code = m.group(2)
        key = f'XCODEX{len(code_blocks)}X'
        try:
            if lang:
                lexer = get_lexer_by_name(lang.strip(), stripall=True)
            else:
                lexer = guess_lexer(code)
        except:
            lexer = TextLexer()
        formatter = HtmlFormatter(nowrap=False, style='monokai' if theme == 'dark' else 'default', noclasses=True)
        highlighted = highlight(code, lexer, formatter)
        code_blocks[key] = f'<div class="code-block">{highlighted}</div>'
        return key

    md_text = re.sub(r'```(\w*)\s*\n(.*?)```', replace_code, md_text, flags=re.DOTALL)

    # Convert markdown to HTML
    html = markdown.markdown(md_text, extensions=[
        'tables', 'fenced_code', 'toc', 'nl2br', 'sane_lists'
    ])

    # Restore code blocks and mermaid
    for key, val in code_blocks.items():
        html = html.replace(f'<p>{key}</p>', val).replace(key, val)
    for key, val in mermaid_blocks.items():
        html = html.replace(f'<p>{key}</p>', val).replace(key, val)

    # Embed local images as base64
    def embed_img(m):
        attrs_before = m.group(1)
        src = m.group(2)
        attrs_after = m.group(3)
        if src.startswith('data:') or src.startswith('http'):
            return m.group(0)
        # Try resolving relative to workspace
        img_path = os.path.join(WORKSPACE, src)
        if os.path.isfile(img_path):
            import mimetypes as mt
            mime = mt.guess_type(img_path)[0] or 'image/png'
            with open(img_path, 'rb') as f:
                b64 = base64.b64encode(f.read()).decode()
            return f'<img {attrs_before}src="data:{mime};base64,{b64}"{attrs_after}'
        return m.group(0)

    html = re.sub(r'<img\s(.*?)src="([^"]+)"(.*?)/?\s*>', embed_img, html)

    return html

# ==================== THEME CSS ====================
def _get_theme_css(theme, is_paginated=True):
    """Get CSS for the given theme. is_paginated controls page-break rules."""
    # No body padding — Playwright PDF margins handle page spacing
    base_css = """
* { box-sizing: border-box; margin: 0; padding: 0; }
body {
    font-family: 'Noto Serif SC', 'Source Han Serif CN', 'Songti SC', Georgia, 'Times New Roman', 'Noto Color Emoji', 'Apple Color Emoji', 'Segoe UI Emoji', serif;
    font-size: 15px;
    line-height: 1.9;
    color: #333;
}
h1 {
    font-size: 26px; font-weight: 700;
    margin: 0 0 20px; padding-bottom: 12px;
    letter-spacing: 0.5px;
}
h2 {
    font-size: 20px; font-weight: 600;
    margin: 36px 0 14px; padding-bottom: 8px;
    letter-spacing: 0.3px;
}
h3 { font-size: 17px; font-weight: 600; margin: 28px 0 10px; }
h4 { font-size: 15px; font-weight: 600; margin: 22px 0 8px; }
h5, h6 { font-size: 14px; font-weight: 600; margin: 16px 0 6px; }
p { margin: 10px 0; text-align: justify; }
ul, ol { margin: 10px 0; padding-left: 24px; }
li { margin: 5px 0; }
blockquote {
    margin: 18px 0;
    padding: 14px 20px 14px 24px;
    border-radius: 4px;
    font-style: normal;
    position: relative;
}
blockquote p { margin: 6px 0; }
table { border-collapse: collapse; margin: 16px 0; width: 100%; font-size: 14px; }
th, td { padding: 10px 14px; text-align: left; }
img { max-width: 100%; max-height: 400px; height: auto; object-fit: contain; border-radius: 6px; margin: 14px auto; display: block; }
hr { margin: 28px 0; border: none; height: 1px; }
a { text-decoration: none; }
.code-block { margin: 16px 0; border-radius: 6px; overflow: hidden; }
.code-block .highlight { padding: 14px 16px; overflow-x: auto; }
.code-block .highlight pre { margin: 0; font-family: 'JetBrains Mono', 'Fira Code', Menlo, Consolas, 'Noto Color Emoji', monospace; font-size: 12.5px; line-height: 1.6; white-space: pre-wrap; word-wrap: break-word; }
code { font-family: 'JetBrains Mono', 'Fira Code', Menlo, Consolas, 'Noto Color Emoji', monospace; font-size: 0.88em; padding: 2px 5px; border-radius: 3px; }
.mermaid-container { margin: 20px 0; text-align: center; }
.mermaid-container svg { max-width: 100%; max-height: 500px; height: auto; }
.mermaid-fallback { padding: 14px; border-radius: 6px; overflow-x: auto; font-size: 13px; }
.signature-block { margin-top: 48px; padding-top: 20px; display: flex; align-items: center; gap: 14px; }
.signature-avatar { width: 44px; height: 44px; border-radius: 50%; object-fit: cover; }
.signature-info { font-size: 13px; }
.signature-name { font-weight: 600; font-size: 14px; }
.signature-date { opacity: 0.5; font-size: 12px; margin-top: 2px; }
"""

    # Pagination CSS — avoid orphan headings, keep blocks together
    if is_paginated:
        base_css += """
h1, h2, h3, h4, h5, h6 { page-break-after: avoid; }
p, li, blockquote { orphans: 3; widows: 3; }
img { page-break-inside: avoid; page-break-before: auto; }
.code-block { page-break-inside: avoid; }
.mermaid-container { page-break-inside: avoid; }
.signature-block { page-break-inside: avoid; }
table { page-break-inside: auto; }
tr { page-break-inside: avoid; }
"""

    themes = {
        'dark': """
body { background: #0d1117; color: #c9d1d9; }
h1 { border-bottom: 1px solid #30363d; color: #f0f6fc; }
h2 { border-bottom: 1px solid #21262d; color: #e6edf3; }
h3, h4, h5, h6 { color: #e6edf3; }
blockquote { background: #161b22; border-left: 4px solid #f0883e; color: #b1bac4; }
th { background: #161b22; border: 1px solid #30363d; color: #f0f6fc; }
td { border: 1px solid #21262d; }
tr:nth-child(even) td { background: rgba(22,27,34,0.5); }
hr { background: #30363d; }
a { color: #58a6ff; }
code { background: #161b22; color: #f0883e; }
.code-block { background: #161b22; border: 1px solid #30363d; }
.mermaid-fallback { background: #161b22; color: #b1bac4; border: 1px solid #30363d; }
.signature-block { border-top: 1px solid #30363d; }
""",
        'sepia': """
body { background: #faf6ed; color: #3d3224; }
h1 { border-bottom: 1px solid #d8cbaf; color: #2a1e0e; }
h2 { border-bottom: 1px solid #ddd0b8; color: #332615; }
h3, h4, h5, h6 { color: #3d2b1a; }
blockquote { background: #f0e8d6; border-left: 4px solid #b8860b; color: #5c4a32; }
th { background: #f0e8d6; border: 1px solid #d8cbaf; color: #2a1e0e; }
td { border: 1px solid #d8cbaf; }
tr:nth-child(even) td { background: #f5efe2; }
hr { background: #d8cbaf; }
a { color: #8b5e3c; }
code { background: #f0e8d6; color: #8b4513; }
.code-block { background: #f0e8d6; border: 1px solid #d8cbaf; }
.mermaid-fallback { background: #f0e8d6; color: #5c4a32; border: 1px solid #d8cbaf; }
.signature-block { border-top: 1px solid #d8cbaf; }
""",
        'light': """
body { background: #ffffff; color: #24292f; }
h1 { border-bottom: 1px solid #d8dee4; color: #1f2328; }
h2 { border-bottom: 1px solid #d8dee4; color: #1f2328; }
h3, h4, h5, h6 { color: #1f2328; }
blockquote { background: #f6f8fa; border-left: 4px solid #0969da; color: #57606a; }
th { background: #f6f8fa; border: 1px solid #d8dee4; color: #1f2328; }
td { border: 1px solid #d8dee4; }
tr:nth-child(even) td { background: #f6f8fa; }
hr { background: #d8dee4; }
a { color: #0969da; }
code { background: #eff1f3; color: #cf222e; }
.code-block { background: #f6f8fa; border: 1px solid #d8dee4; }
.mermaid-fallback { background: #f6f8fa; color: #57606a; border: 1px solid #d8dee4; }
.signature-block { border-top: 1px solid #d8dee4; }
""",
    }
    return base_css + themes.get(theme, themes['light'])

def _signature_html(sig_key, custom_sig=None):
    """Generate signature HTML block. custom_sig: {name, avatar (base64 data URL or None)}."""
    if sig_key == 'custom' and custom_sig:
        name = custom_sig.get('name', '').strip()
        if not name:
            return ''
        avatar_html = ''
        avatar_b64 = custom_sig.get('avatar')
        if avatar_b64:
            # Already a data URL (data:image/...;base64,...) — use directly
            avatar_html = f'<img class="signature-avatar" src="{avatar_b64}" alt="">'
        date_str = datetime.now().strftime('%Y-%m-%d')
        from html import escape
        return f'''
    <div class="signature-block">
        {avatar_html}
        <div class="signature-info">
            <div class="signature-name">{escape(name)}</div>
            <div class="signature-date">{date_str}</div>
        </div>
    </div>'''
    sig = SIGNATURES.get(sig_key)
    if not sig:
        return ''
    avatar_html = ''
    if sig['avatar'] and os.path.isfile(sig['avatar']):
        with open(sig['avatar'], 'rb') as f:
            b64 = base64.b64encode(f.read()).decode()
        avatar_html = f'<img class="signature-avatar" src="data:image/jpeg;base64,{b64}" alt="">'
    date_str = datetime.now().strftime('%Y-%m-%d')
    return f'''
    <div class="signature-block">
        {avatar_html}
        <div class="signature-info">
            <div class="signature-name">{sig["name"]}</div>
            <div class="signature-date">{date_str}</div>
        </div>
    </div>'''

# ==================== PDF EXPORT ====================
# Page layout presets: {name: (width, height, margins)}
PAGE_LAYOUTS = {
    'a4':     {'width': '210mm', 'height': '297mm', 'margin': {'top': '24mm', 'bottom': '24mm', 'left': '20mm', 'right': '20mm'}},
    'a4-compact': {'width': '210mm', 'height': '297mm', 'margin': {'top': '16mm', 'bottom': '16mm', 'left': '16mm', 'right': '16mm'}},
    'letter': {'width': '8.5in', 'height': '11in', 'margin': {'top': '1in', 'bottom': '1in', 'left': '0.8in', 'right': '0.8in'}},
    'a3':     {'width': '297mm', 'height': '420mm', 'margin': {'top': '24mm', 'bottom': '24mm', 'left': '24mm', 'right': '24mm'}},
    'single': None,  # single continuous page, no pagination
}

def export_pdf(content, theme='dark', signature='none', title='document', layout='a4', custom_sig=None):
    """Export markdown to PDF. Returns (filepath, filename, export_id) or raises."""
    is_paginated = layout != 'single'
    html_body = _md_to_html(content, theme)
    css = _get_theme_css(theme, is_paginated=is_paginated)
    sig_html = _signature_html(signature, custom_sig=custom_sig)

    # For single-page, add padding via CSS since no PDF margins
    body_padding = 'padding: 48px 56px;' if not is_paginated else ''

    full_html = f'''<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<style>{css}
body {{ {body_padding} }}
</style>
</head><body>
{html_body}
{sig_html}
</body></html>'''

    export_id = str(uuid.uuid4())[:12]
    safe_title = re.sub(r'[^\w\u4e00-\u9fff\-. ]', '_', title)[:80] or 'document'
    filename = f'{safe_title}.pdf'
    out_path = f'/tmp/md-export-{export_id}.pdf'

    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        # Set viewport to match target PDF width for accurate height measurement
        page.set_viewport_size({'width': 794, 'height': 600})  # 210mm ≈ 794px
        page.set_content(full_html, wait_until='networkidle')
        page.wait_for_timeout(800)

        if not is_paginated:
            # Single continuous page — measure actual rendered height
            height = page.evaluate('''() => {
                const body = document.body;
                const html = document.documentElement;
                return Math.max(body.scrollHeight, body.offsetHeight,
                    html.clientHeight, html.scrollHeight, html.offsetHeight);
            }''')
            # Add generous buffer to prevent overflow to second page
            page.pdf(
                path=out_path,
                width='210mm',
                height=f'{height + 120}px',
                print_background=True,
                margin={'top': '0', 'bottom': '0', 'left': '0', 'right': '0'},
            )
        else:
            # Paginated layout with proper margins
            pl = PAGE_LAYOUTS.get(layout, PAGE_LAYOUTS['a4'])
            page.pdf(
                path=out_path,
                width=pl['width'],
                height=pl['height'],
                print_background=True,
                margin=pl['margin'],
            )

        browser.close()

    return out_path, filename, export_id

# ==================== WORD EXPORT ====================
# Default Word style presets (used when config.json is missing or incomplete)
_DEFAULT_WORD_STYLES = {
    'classic': {
        'font': 'Times New Roman', 'font_east': '宋体', 'size': Pt(12), 'heading_sizes': [Pt(22), Pt(18), Pt(15), Pt(13)],
        'line_spacing': 1.5, 'color': (0x1f, 0x1f, 0x1f)},
    'modern': {
        'font': 'Calibri', 'font_east': '微软雅黑', 'size': Pt(11), 'heading_sizes': [Pt(24), Pt(20), Pt(16), Pt(14)],
        'line_spacing': 1.15, 'color': (0x24, 0x29, 0x2f)},
    'elegant': {
        'font': 'Georgia', 'font_east': '楷体', 'size': Pt(11.5), 'heading_sizes': [Pt(22), Pt(18), Pt(15), Pt(13)],
        'line_spacing': 1.6, 'color': (0x33, 0x33, 0x33)},
}

_DEFAULT_WORD_MARGINS = {
    'normal':  {'top': Cm(2.54), 'bottom': Cm(2.54), 'left': Cm(3.18), 'right': Cm(3.18)},
    'narrow':  {'top': Cm(1.27), 'bottom': Cm(1.27), 'left': Cm(1.27), 'right': Cm(1.27)},
    'wide':    {'top': Cm(2.54), 'bottom': Cm(2.54), 'left': Cm(5.08), 'right': Cm(5.08)},
    'compact': {'top': Cm(1.9),  'bottom': Cm(1.9),  'left': Cm(2.0),  'right': Cm(2.0)},
}

_DEFAULT_WORD_PAGES = {
    'a4':     {'width': Cm(21.0), 'height': Cm(29.7)},
    'letter': {'width': Cm(21.59), 'height': Cm(27.94)},
}

# Heading size defaults per font size range
_HEADING_SIZE_MAP = {
    'small':  [Pt(22), Pt(18), Pt(15), Pt(13)],  # size <= 11.5
    'medium': [Pt(24), Pt(20), Pt(16), Pt(14)],  # size > 11.5
}

def _get_word_styles():
    """Get word styles from config, converting JSON values to docx types."""
    config = get_config()
    cfg_styles = config.get('wordStyles', None)
    if cfg_styles is None:
        return _DEFAULT_WORD_STYLES
    result = {}
    for key, val in cfg_styles.items():
        size_val = val.get('size', 11)
        heading_sizes = _HEADING_SIZE_MAP['medium'] if size_val > 11.5 else _HEADING_SIZE_MAP['small']
        result[key] = {
            'font': val.get('font', 'Calibri'),
            'font_east': val.get('fontEast', '微软雅黑'),
            'size': Pt(size_val),
            'heading_sizes': heading_sizes,
            'line_spacing': val.get('lineSpacing', 1.15),
            'color': (0x24, 0x29, 0x2f),  # default dark color
        }
    return result

def _get_word_margins():
    """Get word margins from config, converting cm floats to Cm()."""
    config = get_config()
    cfg_margins = config.get('wordMargins', None)
    if cfg_margins is None:
        return _DEFAULT_WORD_MARGINS
    result = {}
    for key, val in cfg_margins.items():
        result[key] = {
            'top': Cm(val.get('top', 2.54)),
            'bottom': Cm(val.get('bottom', 2.54)),
            'left': Cm(val.get('left', 3.18)),
            'right': Cm(val.get('right', 3.18)),
        }
    return result

def _get_word_pages():
    """Get word page sizes from config, converting cm floats to Cm()."""
    config = get_config()
    cfg_pages = config.get('wordPages', None)
    if cfg_pages is None:
        return _DEFAULT_WORD_PAGES
    result = {}
    for key, val in cfg_pages.items():
        result[key] = {
            'width': Cm(val.get('width', 21.0)),
            'height': Cm(val.get('height', 29.7)),
        }
    return result

# Backward-compatible module-level references
WORD_STYLES = _DEFAULT_WORD_STYLES
WORD_MARGINS = _DEFAULT_WORD_MARGINS
WORD_PAGES = _DEFAULT_WORD_PAGES

def export_word(content, signature='none', title='document', word_style='modern', word_margin='normal', word_page='a4', custom_sig=None):
    """Export markdown to Word .docx. Returns (filepath, filename) or raises."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    import io

    # Read from config (hot-reloadable)
    word_styles = _get_word_styles()
    word_margins = _get_word_margins()
    word_pages = _get_word_pages()
    ws = word_styles.get(word_style, word_styles.get('modern', _DEFAULT_WORD_STYLES['modern']))
    wm = word_margins.get(word_margin, word_margins.get('normal', _DEFAULT_WORD_MARGINS['normal']))
    wp = word_pages.get(word_page, word_pages.get('a4', _DEFAULT_WORD_PAGES['a4']))

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = ws['font']
    font.size = ws['size']
    font.color.rgb = RGBColor(*ws['color'])
    style.element.rPr.rFonts.set(qn('w:eastAsia'), ws['font_east'])
    # Default line spacing
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = ws['line_spacing']

    # Configure heading styles
    heading_colors = [(0x1a, 0x1a, 0x2e), (0x24, 0x29, 0x2f), (0x33, 0x33, 0x40), (0x44, 0x44, 0x55)]
    for i, hsize in enumerate(ws['heading_sizes'], 1):
        try:
            hs = doc.styles[f'Heading {i}']
            hs.font.size = hsize
            hs.font.color.rgb = RGBColor(*heading_colors[min(i-1, 3)])
            hs.font.bold = True
            hs.font.name = ws['font']
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), ws['font_east'])
        except:
            pass

    # Pre-process mermaid blocks → PNG images (parallel)
    mermaid_w_codes = []
    mermaid_w_keys = []
    def extract_mermaid_w(m):
        code = m.group(1).strip()
        key = f'XMERMAIDWX{len(mermaid_w_codes)}X'
        mermaid_w_codes.append(code)
        mermaid_w_keys.append(key)
        return key
    content = re.sub(r'```mermaid\s*\n(.*?)```', extract_mermaid_w, content, flags=re.DOTALL)

    mermaid_images = {}
    if mermaid_w_codes:
        pngs = _render_mermaid_pngs_batch(mermaid_w_codes)
        for key, png_data in zip(mermaid_w_keys, pngs):
            if png_data:
                mermaid_images[key] = png_data

    # Set page size and margins
    for section in doc.sections:
        section.page_width = wp['width']
        section.page_height = wp['height']
        section.top_margin = wm['top']
        section.bottom_margin = wm['bottom']
        section.left_margin = wm['left']
        section.right_margin = wm['right']

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lang = ''
    code_lines = []
    in_list = False

    def add_inline_formatting(paragraph, text):
        """Parse inline markdown (bold, italic, code, links) and add runs."""
        # Pattern for **bold**, *italic*, `code`, [text](url)
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[([^\]]+)\]\(([^)]+)\))'
        pos = 0
        for m in re.finditer(pattern, text):
            # Add text before match
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            if m.group(2):  # bold
                run = paragraph.add_run(m.group(2))
                run.bold = True
            elif m.group(3):  # italic
                run = paragraph.add_run(m.group(3))
                run.italic = True
            elif m.group(4):  # code
                run = paragraph.add_run(m.group(4))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xcf, 0x22, 0x2e)
            elif m.group(5):  # link
                run = paragraph.add_run(m.group(5))
                run.font.color.rgb = RGBColor(0x09, 0x69, 0xda)
                run.underline = True
            pos = m.end()
        # Remaining text
        if pos < len(text):
            paragraph.add_run(text[pos:])

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x1f, 0x23, 0x28)
                # Add gray background via shading
                from docx.oxml import OxmlElement
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F6F8FA')
                shd.set(qn('w:val'), 'clear')
                p.paragraph_format.element.get_or_add_pPr().append(shd)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lang = line[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Mermaid placeholder
        for key, png_data in mermaid_images.items():
            if key in line:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(io.BytesIO(png_data), width=Inches(5))
                i += 1
                continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            heading_level = min(level, 4)
            h = doc.add_heading(text, level=heading_level)
            # Keep heading with next paragraph (no orphan headings)
            h.paragraph_format.keep_with_next = True
            # Add spacing before H1/H2 for visual separation
            if level <= 2:
                h.paragraph_format.space_before = Pt(24)
                h.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # Blockquote
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x57, 0x60, 0x6a)
            p.paragraph_format.left_indent = Cm(1.5)
            # Add left border via shading
            from docx.oxml import OxmlElement
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:color'), '0969DA')
            left.set(qn('w:space'), '8')
            pBdr.append(left)
            p.paragraph_format.element.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', line.strip()):
            p = doc.add_paragraph()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), 'D1D9E0')
            bottom.set(qn('w:space'), '1')
            pBdr.append(bottom)
            p.paragraph_format.element.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r'^(\s*)[-*+]\s+(.+)', line)
        if ul_match:
            indent = len(ul_match.group(1))
            text = ul_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            add_inline_formatting(p, text)
            if indent >= 2:
                p.paragraph_format.left_indent = Cm(1.5 + indent * 0.3)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if ol_match:
            indent = len(ol_match.group(1))
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_inline_formatting(p, text)
            if indent >= 2:
                p.paragraph_format.left_indent = Cm(1.5 + indent * 0.3)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline_formatting(p, line)
        i += 1

    # Signature
    sig = None
    custom_name = None
    custom_avatar_b64 = None
    if signature == 'custom' and custom_sig:
        custom_name = custom_sig.get('name', '').strip()
        custom_avatar_b64 = custom_sig.get('avatar')  # data URL or None
        if custom_name:
            sig = True  # flag to render custom
    else:
        sig = SIGNATURES.get(signature)

    if sig:
        doc.add_paragraph()  # spacer
        p = doc.add_paragraph()
        date_str = datetime.now().strftime('%Y-%m-%d')
        if signature == 'custom' and custom_name:
            # Custom signature: decode base64 avatar if present
            if custom_avatar_b64:
                try:
                    # Strip data URL prefix: "data:image/png;base64,..."
                    if ',' in custom_avatar_b64:
                        img_bytes = base64.b64decode(custom_avatar_b64.split(',', 1)[1])
                    else:
                        img_bytes = base64.b64decode(custom_avatar_b64)
                    run = p.add_run()
                    run.add_picture(io.BytesIO(img_bytes), width=Inches(0.5))
                    p.add_run('  ')
                except Exception:
                    pass  # skip avatar on decode failure
            run = p.add_run(custom_name)
            run.bold = True
            run.font.size = Pt(12)
            p.add_run(f'    {date_str}')
        else:
            # Preset signature
            if sig['avatar'] and os.path.isfile(sig['avatar']):
                run = p.add_run()
                run.add_picture(sig['avatar'], width=Inches(0.5))
                p.add_run('  ')
            run = p.add_run(sig['name'])
            run.bold = True
            run.font.size = Pt(12)
            p.add_run(f'    {date_str}')

    # Add page numbers in footer
    try:
        from docx.oxml import OxmlElement
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Page number field
        run = fp.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
    except:
        pass

    export_id = str(uuid.uuid4())[:12]
    safe_title = re.sub(r'[^\w\u4e00-\u9fff\-. ]', '_', title)[:80] or 'document'
    filename = f'{safe_title}.docx'
    out_path = f'/tmp/md-export-{export_id}.docx'
    doc.save(out_path)

    return out_path, filename, export_id
